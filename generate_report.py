"""
generate_report.py
==================
Generates the ECO-FAST Tool Description Report as a Word document.
Run with: py -3.12 generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.18)
    section.right_margin  = Cm(3.18)

# ── Helpers ───────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 73, 125)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(8)
    p.paragraph_format.space_before      = Pt(0)
    if indent == 0:
        p.paragraph_format.first_line_indent = Cm(0.75)
    else:
        p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    set_font(run)
    return p

def bullet(label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    r1 = p.add_run(label)
    set_font(r1, bold=True)
    r2 = p.add_run(text)
    set_font(r2)
    return p

def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    cell._tc.get_or_add_tcPr().append(shd)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, "003366")
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=11, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        fill = "EAF0FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            shade_cell(cell, fill)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=11)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ECO-FAST")
set_font(r, size=36, bold=True, color=(0, 51, 102))

doc.add_paragraph()

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run(
    "Economically and Environmentally Conscious\n"
    "Food Waste Assessment and Selection Tool"
)
set_font(r, size=18, italic=True, color=(31, 73, 125))

doc.add_paragraph()
doc.add_paragraph()

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run("TOOL DESCRIPTION REPORT")
set_font(r, size=14, bold=True)

doc.add_paragraph()

dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dt.add_run(f"Version 1.0   \u2014   {datetime.date.today().strftime('%B %Y')}")
set_font(r, size=12, italic=True, color=(89, 89, 89))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
heading1("Table of Contents")
toc = [
    ("1.", "Overview and Purpose"),
    ("2.", "System Architecture"),
    ("3.", "Technology Superstructure"),
    ("4.", "Tab 1 \u2014 Instructions"),
    ("5.", "Tab 2 \u2014 Feed Inputs"),
    ("6.", "Tab 3 \u2014 Technology Specifications"),
    ("7.", "Tab 4 \u2014 Cost Specifications"),
    ("8.", "Tab 5 \u2014 Results"),
    ("9.", "Tab 6 \u2014 Environmental Justice Assessment"),
    ("10.", "Mathematical Model Summary"),
    ("11.", "Data Flow Between Tabs"),
    ("12.", "Software Stack and Deployment"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{num}  ")
    set_font(r1, bold=True)
    r2 = p.add_run(title)
    set_font(r2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. OVERVIEW AND PURPOSE
# ═══════════════════════════════════════════════════════════════
heading1("1.  Overview and Purpose")

body(
    "ECO-FAST (Economically and Environmentally Conscious Food Waste Assessment "
    "and Selection Tool) is a web-based decision-support platform developed to "
    "assist engineers, municipalities, waste management authorities, and "
    "policymakers in identifying the most economically viable and environmentally "
    "responsible approach for processing a given food waste stream. The tool is "
    "designed to be accessible to users without deep optimisation expertise, "
    "while providing rigorous, mathematically grounded results suitable for "
    "academic publication and regulatory decision-making."
)

body(
    "Given a characterised food waste stream \u2014 defined by its flow rate, "
    "wet-basis composition, and facility location \u2014 ECO-FAST evaluates all "
    "feasible combinations of pretreatment, conversion, and product recovery "
    "technologies simultaneously using a Mixed-Integer Nonlinear Programming "
    "(MINLP) optimisation model. The model selects the optimal technology "
    "pathway and computes the associated annualised costs and greenhouse gas "
    "(GHG) emissions for three objective modes:"
)

bullet("Lowest Cost Pathway:  ",
    "Identifies the processing configuration that minimises Net Annual Cost "
    "(NAC, M$ yr\u207b\u00b9) without any constraint on GHG emissions. "
    "This represents the economically dominant solution.")
bullet("Lowest Emissions Pathway:  ",
    "Identifies the configuration that minimises total life-cycle GHG emissions "
    "(t CO\u2082-eq yr\u207b\u00b9) without constraint on cost. "
    "This represents the environmentally dominant solution.")
bullet("Lowest Cost and Emissions Pathway:  ",
    "Applies the epsilon-constraint multi-objective method to generate the full "
    "Pareto frontier of cost-versus-emissions trade-offs. Each point on the "
    "frontier is Pareto-optimal \u2014 no other feasible configuration can improve "
    "one objective without worsening the other. This mode supports decision-making "
    "where both economic and environmental performance must be balanced.")

body(
    "The tool is intended for use at the facility-planning stage, where a waste "
    "management authority or developer is evaluating which processing technology "
    "to invest in. Results are expressed in annualised terms at the design "
    "throughput. Sensitivity to economic assumptions (prices, discount rate, "
    "plant lifetime) can be explored by adjusting parameters in Tabs 3 and 4 "
    "and re-running the optimisation."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
heading1("2.  System Architecture")

body(
    "ECO-FAST follows a two-process client-server architecture. Both processes "
    "must run simultaneously on the same machine for the tool to function correctly."
)

heading2("2.1  Frontend \u2014 Streamlit (Port 8501)")
body(
    "The user-facing web application is built with Streamlit and served at "
    "http://localhost:8501. It renders all six tabs, collects user inputs "
    "through interactive widgets, sends optimisation requests to the backend "
    "as JSON payloads via the FastAPI REST API, and displays results including "
    "tables, charts, and interactive maps. All inputs entered during a session "
    "are held in Streamlit session state and are lost when the browser is closed "
    "or the page is refreshed."
)

heading2("2.2  Backend \u2014 FastAPI (Port 8000)")
body(
    "A REST API server built with FastAPI and served by Uvicorn at "
    "http://127.0.0.1:8000. It receives optimisation requests from the frontend, "
    "builds the full GAMSPy MINLP model from the provided inputs, executes the "
    "solver, and returns results as a structured JSON response. The backend is "
    "stateless; each request builds and solves a fresh model instance."
)

heading2("2.3  Optimisation Model Package (model/)")
body(
    "The mathematical model is implemented as a Python package under the model/ "
    "directory comprising seven modules:"
)
bullet("sets.py:  ", "Defines all index sets (technologies, stream components, product types).")
bullet("parameters.py:  ", "Declares all scalar and indexed parameters with literature-based default values.")
bullet("variables.py:  ", "Declares all decision variables (binary technology selectors, continuous mass flows, economic and emission quantities).")
bullet("equations.py:  ", "Implements all mass balance, process constraint, and sizing equations across all 15 technology units.")
bullet("costing.py:  ", "Computes all eight cost components (capital, working capital, insurance, utility, labour, overhead, raw material, disposal) and total revenue.")
bullet("ghg.py:  ", "Computes all GHG emission terms (indirect, direct, aqueous-phase, displacement credits).")
bullet("solver.py:  ", "Implements the three objective modes, the epsilon-constraint Pareto logic, and the run_optimization() entry point called by the API.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. TECHNOLOGY SUPERSTRUCTURE
# ═══════════════════════════════════════════════════════════════
heading1("3.  Technology Superstructure")

body(
    "The superstructure represents all possible processing pathways considered "
    "by the optimisation model. It is organised into four sequential stages "
    "connected by 61 numbered material streams. Each stream carries "
    "component-resolved mass flow rates (kg hr\u207b\u00b9) across eight "
    "components: water (WATER), carbohydrates (CBH), protein (PRT), fat/lipid "
    "(FAT), other organics (OTH), ash (ASH), fixed carbon (FC), and enzyme "
    "(ENZYME). Binary decision variables select one configuration at each stage, "
    "enforcing mutual exclusivity through equality constraints."
)

heading2("3.1  Stage 1 \u2014 Mechanical Pretreatment")
body(
    "One of three mechanical pretreatment options is selected by binary variable "
    "y_{s1,m} in {0,1}:"
)
bullet("Shredding (SHR):  ",
    "Dry particle-size reduction. Specific power 0.05 kW per kg hr\u207b\u00b9 "
    "throughput. Suitable for all downstream conversion pathways.")
bullet("Maceration (MCR):  ",
    "Wet grinding with water addition (default ratio 2.0 kg water per kg dry "
    "feed) producing a pumpable slurry. Used primarily ahead of HTL and AND.")
bullet("Bypass (BYP):  ",
    "Feed passes directly to biological pretreatment or conversion without "
    "mechanical treatment. Required for landfill and incineration pathways.")

heading2("3.2  Stage 2 \u2014 Biological Pretreatment")
body(
    "An optional biological pretreatment step conditions the feed prior to "
    "conversion. One of three options is selected:"
)
bullet("Aerobic Biodigestion (AER):  ",
    "Partial aerobic degradation of volatile solids fractions at "
    "component-specific degradation efficiencies (CBH 0.40, PRT 0.30, "
    "FAT 0.15, OTH 0.20 by default). Produces a CO\u2082/H\u2082O offgas "
    "and a conditioned slurry fed to conversion.")
bullet("Enzymatic Hydrolysis (ENZ):  ",
    "Cellulase enzyme addition (default dose 0.02 kg enzyme per kg dry solids) "
    "that hydrolyses carbohydrates into fermentable sugars, improving biogas "
    "yields in downstream AND. Hydrolysis efficiency default 0.85.")
bullet("Bypass (BYP_bio):  ",
    "Feed passes directly to the conversion stage. Required for CMP, SLF, "
    "INC, and WWT pathways.")

heading2("3.3  Stage 3 \u2014 Conversion")
body(
    "The primary decision of the model: one of six mutually exclusive "
    "conversion technologies processes the conditioned feed."
)

add_table(
    ["Technology", "Abbreviation", "Primary Product", "Key Parameters"],
    [
        ["Hydrothermal Liquefaction", "HTL", "Biocrude oil",
         "Temperature 300-350 C; pressure 150-250 bar; biocrude yield 0.35 (default)"],
        ["Anaerobic Digestion", "AND", "Biomethane",
         "BMP 400 mL CH4/g VS; VS degradation efficiency 0.85; HRT 720 hr"],
        ["Composting", "CMP", "Compost",
         "VS degradation 50%; composting time 120 hr; IPCC EFs for CH4 and N2O"],
        ["Sanitary Landfill", "SLF", "Landfill gas (CH4)",
         "IPCC first-order decay; gas capture efficiency 0.65; MCF 0.60"],
        ["Wastewater Treatment", "WWT", "Biosolids",
         "F/M ratio 0.30 kg BOD/kg MLSS/day; MLSS 3,000 mg/L; SRT 10 days"],
        ["Incineration", "INC", "Electricity (via STB)",
         "Excess air ratio lambda=1.2; elemental combustion stoichiometry; coupled to steam turbine"],
    ],
    col_widths=[4.0, 3.0, 3.5, 5.5]
)

heading2("3.4  Stage 4 \u2014 Product Recovery and Upgrading")
body(
    "Downstream units recover and upgrade conversion products to marketable form:"
)
bullet("Centrifugation (CEN) / Filtration (FLT):  ",
    "Separate biocrude, char, and aqueous phases from the HTL product stream. "
    "CEN achieves 95% biocrude recovery; FLT achieves 98%. One unit is selected "
    "when HTL is the chosen conversion pathway.")
bullet("Amine Scrubbing (ABS):  ",
    "Chemical absorption of CO\u2082 from biogas to produce pipeline-quality "
    "biomethane (>=96% CH\u2084 purity). CH\u2084 slip 0.1%, CO\u2082 "
    "removal 98.5%, electricity use 0.12 kWh m\u207b\u00b3.")
bullet("Pressure Swing Adsorption (PSA):  ",
    "Physical adsorption of CO\u2082 from biogas. CH\u2084 slip 1.5%, "
    "CO\u2082 removal 97%, electricity use 0.25 kWh m\u207b\u00b3.")
bullet("Steam Turbine (STB):  ",
    "Recovers energy from INC flue gas to raise superheated steam "
    "(h1 = 3,277.9 kJ/kg) that expands through a back-pressure turbine "
    "(eta_turbine = 0.85, eta_generator = 0.95, h2 = 2,800 kJ/kg) to generate electricity.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. TAB 1 — INSTRUCTIONS
# ═══════════════════════════════════════════════════════════════
heading1("4.  Tab 1 \u2014 Instructions")

body(
    "The Instructions tab is the landing page of ECO-FAST. It requires no user "
    "input and is intended to orient new users before they begin entering data. "
    "The tab contains the following content:"
)

bullet("Tool purpose statement:  ",
    "A concise description of what ECO-FAST does and the three objective modes "
    "available (Lowest Cost, Lowest Emissions, Lowest Cost and Emissions).")
bullet("Superstructure diagram:  ",
    "A graphical representation of the full technology network (loaded from "
    "Superstructure.tif). The diagram shows all 15 processing units and their "
    "interconnections across all four processing stages.")
bullet("Technology summary panels:  ",
    "Three side-by-side panels listing the technologies available at each stage: "
    "Pretreatment (SHR, MCR, AER, ENZ), Conversion (HTL, AND, CMP, SLF, WWT, INC), "
    "and Recovery and Upgrading (CEN, FLT, ABS, PSA, STB).")
bullet("Step-by-step usage guide:  ",
    "Five numbered instructions directing the user through Tabs 2 to 6 in sequence.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. TAB 2 — FEED INPUTS
# ═══════════════════════════════════════════════════════════════
heading1("5.  Tab 2 \u2014 Feed Inputs")

body(
    "The Feed Inputs tab is the primary data entry interface where the user "
    "characterises the food waste stream entering the system. All values entered "
    "here are stored in session state and passed to the optimisation backend at "
    "the time of execution. The tab is divided into a left input panel and a "
    "right visualisation panel."
)

heading2("5.1  Food Waste Type")
body(
    "The user selects a food waste category from a predefined library. The "
    "selection automatically populates all composition fields with representative "
    "wet-basis mass fractions from the published literature. A 'Custom (Enter "
    "Your Own Data)' option unlocks all seven composition fields for manual entry. "
    "The tool does not allow the user to proceed further until a waste type is selected."
)

heading2("5.2  Feed Stream Conditions")
add_table(
    ["Input", "Unit Options", "Default", "Description"],
    [
        ["Feed flow rate", "kg/hr, t/day, lb/hr, kg/day", "1,000 kg/hr",
         "Mass flow rate of the wet food waste feed stream entering the system. Converted internally to kg/hr for all calculations."],
        ["Annual operating hours", "hr/yr, days/yr, weeks/yr", "7,920 hr/yr",
         "Total hours per year the facility operates. Default corresponds to 330 days/yr of continuous operation. Converted internally to hr/yr."],
    ],
    col_widths=[3.5, 3.5, 2.5, 5.5]
)

heading2("5.3  Facility Location")
body(
    "The user enters a five-digit US zip code identifying the proposed facility "
    "location. This field is optional for the optimisation but required for the "
    "Environmental Justice assessment in Tab 6. The zip code is used to retrieve "
    "geographic coordinates from the Zippopotam.us API and demographic data from "
    "the US Census Bureau ACS API. If no zip code is entered, Tab 6 will remain inactive."
)

heading2("5.4  Wet-Basis Composition")
body(
    "Seven component mass fractions are entered on a wet basis. When a predefined "
    "waste type is selected, these fields are auto-populated and locked. For the "
    "custom option, all fields are editable. A running total is displayed; the "
    "tool raises an error and blocks execution if the sum deviates from 1.0 by "
    "more than 0.01."
)
add_table(
    ["Component", "Symbol", "Typical Range (wet basis)", "Role in the Model"],
    [
        ["Water", "WATER", "0.60 - 0.85",
         "Determines drying and evaporation loads. Dilutes organic content, reducing volumetric energy density and biogas yield per unit feed."],
        ["Carbohydrates", "CBH", "0.01 - 0.12",
         "Primary substrate for biological conversion pathways. Governs BMP contribution from carbohydrates and AER/ENZ degradation kinetics."],
        ["Protein", "PRT", "0.02 - 0.15",
         "Key nitrogen source. Determines C/N ratio (important for composting stability and AD inhibition risk). Participates in AER and ENZ degradation."],
        ["Fat / Lipid", "FAT", "0.01 - 0.08",
         "Highest energy density of all components. Disproportionately contributes to biocrude yield in HTL and to methane yield in AND due to high BMP of lipids."],
        ["Other Organics", "OTH", "0.01 - 0.05",
         "Catch-all for non-classified organic matter (e.g. cellulose, lignin, mixed organics). Participates in all biological and thermal conversion pathways."],
        ["Ash", "ASH", "0.005 - 0.02",
         "Inert mineral content. Passes through all conversion units unchanged and reports to residue, char, or biosolids output streams."],
        ["Fixed Carbon", "FC", "0.01 - 0.05",
         "Stable carbonaceous material (e.g. charcoal-like fractions). Slow to degrade biologically; contributes to char yield in HTL and increases ash-like residue in INC."],
    ],
    col_widths=[3.0, 2.0, 3.0, 7.0]
)

heading2("5.5  Feedstock Quality Indicators")
body(
    "The right panel automatically computes and displays four quality indicators "
    "from the entered composition. These metrics give the user an immediate "
    "sense of which processing pathways are likely to be feasible and economically "
    "attractive for their specific feedstock."
)
add_table(
    ["Indicator", "Basis", "Interpretation"],
    [
        ["Higher Heating Value (HHV, MJ/kg)",
         "Correlation based on organic component fractions",
         "Gross energy content of the feed. Higher HHV favours INC and HTL pathways for energy recovery."],
        ["Volatile Solids (VS, %)",
         "(CBH + PRT + FAT + OTH) x 100",
         "Biodegradable organic fraction. Primary basis for BMP and degradation calculations in AND, AER, CMP, and WWT."],
        ["Total Solids (TS, %)",
         "(1 - WATER) x 100",
         "Dry matter content. Low TS (high moisture) limits thermochemical conversion efficiency but suits aqueous processes like HTL and WWT."],
        ["C/N Ratio",
         "Carbon in organics / Nitrogen in protein",
         "Indicator of composting suitability (optimal 25-35:1) and AD stability. High C/N may inhibit microbial activity in biological pathways."],
    ],
    col_widths=[4.0, 4.5, 6.5]
)
body(
    "A donut chart in the right panel visualises the wet-basis composition "
    "breakdown by component with a colour-coded legend, updating in real time "
    "as the user changes composition values."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. TAB 3 — TECHNOLOGY SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════
heading1("6.  Tab 3 \u2014 Technology Specifications")

body(
    "The Technology Specifications tab provides access to the engineering "
    "parameters for each of the 15 processing units in the superstructure. "
    "Users can review and adjust default values before running the optimisation. "
    "Default values are drawn from peer-reviewed literature, IPCC 2006 Tier 1 "
    "guidelines, and standard chemical engineering references. A dropdown menu "
    "at the top of the tab selects the technology; the corresponding parameter "
    "panel appears below."
)

heading2("6.1  Hydrothermal Liquefaction (HTL)")
add_table(
    ["Parameter", "Unit", "Default", "Description"],
    [
        ["Biocrude yield fraction", "-", "0.35", "Mass fraction of DAF feed converted to biocrude oil product"],
        ["Char yield fraction", "-", "0.10", "Mass fraction converted to solid char by-product"],
        ["Aqueous phase yield fraction", "-", "0.40", "Mass fraction reporting to the aqueous by-product phase"],
        ["Gas product yield fraction", "-", "0.15", "Mass fraction released as non-condensable gas; sum of four yield fractions must equal 1.0"],
        ["Reaction temperature", "C", "340", "HTL reactor operating temperature; 280-370 C range for food waste"],
        ["Residence time", "min", "60", "Mean residence time in the HTL reactor"],
        ["Dry solid loading", "%", "7.0", "Mass fraction of dry solids in the feed slurry entering the reactor"],
    ],
    col_widths=[5.0, 1.5, 2.0, 6.5]
)

heading2("6.2  Anaerobic Digestion (AND)")
add_table(
    ["Parameter", "Unit", "Default", "Description"],
    [
        ["Biochemical methane potential (BMP)", "mL CH4/g VS", "400", "Maximum specific methane yield under batch test conditions; represents the biodegradability of the VS fraction. Literature range for food waste: 357-490 mL CH4/g VS."],
        ["VS degradation efficiency", "-", "0.85", "Fraction of BMP achieved in a continuous digester; accounts for kinetic limitations relative to the batch BMP assay"],
        ["Biogas capture efficiency", "-", "0.98", "Fraction of produced biogas successfully collected"],
        ["Other organics degradation fraction", "-", "0.30", "Fraction of OTH component degraded in the digester"],
        ["Hydraulic retention time (HRT)", "hr", "720", "Mean residence time of feed in the digester (30 days default)"],
        ["Organic loading rate (OLR)", "kg VS/m3/day", "2.5", "Mass of volatile solids fed per unit reactor volume per day; determines reactor volume"],
        ["Vessel fill fraction", "-", "0.85", "Fraction of the reactor volume available for liquid; accounts for headspace"],
    ],
    col_widths=[5.5, 2.5, 2.0, 5.0]
)

heading2("6.3  Aerobic Biodigestion (AER)")
add_table(
    ["Parameter", "Unit", "Default", "Description"],
    [
        ["Carbohydrate degradation fraction (fdeg_CBH)", "-", "0.40", "Fraction of carbohydrates aerobically degraded to CO2 and H2O"],
        ["Protein degradation fraction (fdeg_PRT)", "-", "0.30", "Fraction of protein degraded aerobically"],
        ["Fat degradation fraction (fdeg_FAT)", "-", "0.15", "Fraction of fat degraded aerobically"],
        ["Other organics degradation (fdeg_OTH)", "-", "0.20", "Fraction of other organics degraded aerobically"],
        ["Hydraulic retention time", "hr", "24", "Residence time in the aerobic biodigester"],
        ["Water addition ratio", "kg/kg dry feed", "1.5", "Water added to maintain slurry consistency for pumping"],
    ],
    col_widths=[6.0, 2.0, 2.0, 5.0]
)

heading2("6.4  Enzymatic Hydrolysis (ENZ)")
add_table(
    ["Parameter", "Unit", "Default", "Description"],
    [
        ["Hydrolysis efficiency (eta_ENZ)", "-", "0.85", "Fraction of carbohydrates successfully hydrolysed to soluble sugars"],
        ["Enzyme dose (r_enz)", "kg enzyme/kg dry solids", "0.02", "Mass of cellulase enzyme added per kg of dry feed solids"],
        ["Hydraulic retention time (HRT_ENZ)", "hr", "6.0", "Contact time between enzyme and substrate"],
    ],
    col_widths=[5.5, 3.5, 2.0, 4.0]
)

heading2("6.5  Composting (CMP)")
add_table(
    ["Parameter", "Unit", "Default", "Source"],
    [
        ["Volatile solids degradation fraction", "-", "0.50", "Standard composting literature"],
        ["Other organics degradation fraction", "-", "0.60", "Process-specific estimate"],
        ["Composting time", "hr", "120", "Typical active composting phase duration (5 days)"],
        ["CH4 emission factor", "kg CH4/kg wet feed", "0.004", "IPCC 2006 Tier 1 default for aerobic composting"],
        ["N2O emission factor", "kg N2O/kg wet feed", "0.0003", "IPCC 2006 Tier 1 default for aerobic composting"],
    ],
    col_widths=[5.5, 2.5, 2.0, 5.0]
)

heading2("6.6  Wastewater Treatment (WWT)")
add_table(
    ["Parameter", "Unit", "Default", "Description"],
    [
        ["BOD fraction of volatile solids (fBOD)", "-", "1.20", "Ratio of biochemical oxygen demand to volatile solids mass; reflects oxygen demand of the organic substrate"],
        ["Sludge retention time (SRT)", "days", "10", "Mean residence time of biological solids in the activated sludge system; governs biosolids production via yield/decay model"],
        ["MLSS concentration (X_MLSS)", "g/L", "3.0", "Mixed liquor suspended solids target concentration; typical activated sludge range 2,000-4,000 mg/L"],
        ["Minimum HRT", "hr", "6.0", "Minimum hydraulic retention time constraint in the aeration basin"],
    ],
    col_widths=[5.5, 1.5, 2.0, 6.0]
)

heading2("6.7  Sanitary Landfill (SLF)")
add_table(
    ["Parameter", "Unit", "Default", "Source"],
    [
        ["Degradable organic carbon fraction (DOC)", "-", "0.358", "IPCC 2006 Tier 1 default for food waste"],
        ["Fraction of DOC that decomposes (DOC_f)", "-", "0.77", "IPCC 2006 Tier 1 default"],
        ["Methane correction factor (MCF)", "-", "0.60", "IPCC 2006 Tier 1 default for managed anaerobic landfill"],
        ["CH4 fraction in landfill gas (F)", "-", "0.576", "IPCC 2006 Tier 1 default"],
        ["Gas capture efficiency", "-", "0.65", "Typical value for landfills with active gas collection"],
        ["Oxidation factor (OX)", "-", "0.10", "IPCC 2006 Tier 1 default; fraction of CH4 oxidised in landfill cover"],
        ["Landfill depth", "m", "10", "Used in area-based sizing calculations"],
    ],
    col_widths=[5.5, 1.5, 2.0, 6.0]
)

heading2("6.8  Incineration (INC)")
body(
    "INC uses a first-principles elemental combustion model rather than a fixed "
    "conversion efficiency. Flue gas composition is computed from mole balances "
    "on C, H, O, N, and S using the dry-basis ultimate analysis of the feed. "
    "The three adjustable parameters are:"
)
add_table(
    ["Parameter", "Unit", "Default", "Description"],
    [
        ["Combustion efficiency", "-", "0.98", "Fraction of combustible material fully oxidised in the furnace"],
        ["Boiler efficiency", "-", "0.80", "Fraction of flue gas thermal energy transferred to the steam circuit"],
        ["Excess air ratio (lambda)", "-", "1.20", "Ratio of actual air supplied to stoichiometric air requirement; 1.20 means 20% excess air"],
    ],
    col_widths=[4.0, 1.5, 2.0, 7.5]
)

heading2("6.9  Gas Upgrading (ABS and PSA)")
add_table(
    ["Parameter", "ABS Default", "PSA Default", "Description"],
    [
        ["CO2 removal efficiency", "98.5%", "97.0%", "Fraction of CO2 removed from the raw biogas stream"],
        ["CH4 slip fraction", "0.1%", "1.5%", "Fraction of CH4 lost to the off-gas (not recovered in product stream)"],
        ["Electricity consumption", "0.12 kWh/m3", "0.25 kWh/m3", "Specific electricity use per m3 of raw biogas processed at standard conditions"],
        ["Required biomethane purity", "96% CH4", "96% CH4", "Minimum molar fraction of CH4 in the product stream (enforced as a model constraint)"],
    ],
    col_widths=[4.0, 2.5, 2.5, 6.0]
)

heading2("6.10  Steam Turbine (STB)")
add_table(
    ["Parameter", "Unit", "Default", "Description"],
    [
        ["Turbine isentropic efficiency (eta_t)", "-", "0.85", "Ratio of actual work extracted to isentropic (ideal) work; represents mechanical losses in the turbine"],
        ["Generator efficiency (eta_g)", "-", "0.95", "Fraction of shaft power converted to electrical power by the generator"],
        ["Turbine inlet enthalpy (h1)", "kJ/kg", "3,277.9", "Specific enthalpy of superheated steam entering the turbine; corresponds to approximately 60 bar and 450 C"],
        ["Turbine outlet enthalpy (h2)", "kJ/kg", "2,800", "Isentropic outlet enthalpy; corresponds to back-pressure exhaust at approximately 9-10 bar"],
        ["Flue gas inlet temperature", "C", "850", "Temperature of INC flue gas entering the waste heat recovery boiler"],
        ["Flue gas outlet temperature", "C", "400", "Temperature of cooled flue gas leaving the boiler"],
        ["Flue gas heat capacity (Cp)", "kJ/kg C", "1.15", "Specific heat of combustion flue gas at elevated temperature"],
    ],
    col_widths=[5.5, 1.5, 2.0, 6.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. TAB 4 — COST SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════
heading1("7.  Tab 4 \u2014 Cost Specifications")

body(
    "The Cost Specifications tab collects the economic parameters and solver "
    "settings required to run the optimisation. It is organised into three "
    "sections: economic parameters, product prices and revenue preview, "
    "and solver configuration."
)

heading2("7.1  Economic Parameters")
add_table(
    ["Parameter", "Unit", "Default", "Role in Model"],
    [
        ["Electricity cost (C_elec)", "$/kWh", "0.10",
         "Unit purchase price of grid electricity. Applied to net electricity consumption (PW_total) in the utility cost equation: CC_UC = C_elec x PW_net x T_ann."],
        ["Tipping fee (C_tip)", "$/kg feed", "0.08",
         "Revenue received per kg of food waste accepted at the gate. Included in the REV calculation: REV_tip = C_tip x Q_f x T_ann / 10^6."],
        ["Labor cost (C_lbr)", "$/hr", "30",
         "Hourly wage rate for facility operators. Labour cost is scaled by the number of workers (N_lbr) determined from facility size correlations."],
        ["Discount rate (r)", "% yr-1", "User defined",
         "Annual discount rate used to compute the Capital Recovery Factor (CRF = r(1+r)^n / ((1+r)^n - 1))."],
        ["Plant lifetime (n)", "yr", "User defined",
         "Economic project lifetime. Together with r, determines CRF and hence the annualised capital cost (CC_AC)."],
    ],
    col_widths=[4.0, 2.0, 2.5, 6.5]
)

heading2("7.2  Product Prices and Revenue Preview")
body(
    "The user sets the market selling price for each saleable product. A live "
    "revenue calculator below the input fields computes and displays the estimated "
    "annual revenue from each product at the current feed rate and composition "
    "(using pre-set yield assumptions), allowing the user to assess the economic "
    "significance of each revenue stream before committing to a full optimisation run."
)
add_table(
    ["Product", "Unit", "Typical Range", "Notes"],
    [
        ["Biocrude", "$/kg", "0.30 - 0.60",
         "Crude bio-oil from HTL pathway. Displaces fossil crude in the GHG displacement credit calculation."],
        ["Biomethane", "$/kg", "0.50 - 1.20",
         "Upgraded biogas (>=96% CH4) from AND or SLF pathways. Displaces fossil natural gas."],
        ["Compost", "$/kg", "0.03 - 0.15",
         "Stabilised organic matter from the CMP pathway. Displaces synthetic fertiliser in GHG accounting."],
        ["Biosolids", "$/kg", "0.01 - 0.10",
         "Dewatered sludge from WWT (and HTL aqueous treatment). Displaces mineral soil amendments."],
        ["Electricity", "$/kWh", "0.08 - 0.15",
         "Generated by the steam turbine in the INC pathway. Displaces grid electricity."],
    ],
    col_widths=[2.5, 1.5, 3.0, 8.0]
)

heading2("7.3  Solver Configuration")
add_table(
    ["Setting", "Options / Range", "Default", "Effect on Computation"],
    [
        ["Optimisation objective", "Lowest Cost / Lowest Emissions / Balanced",
         "Balanced",
         "Selects which of the three solver modes is executed. 'Lowest Cost' and 'Lowest Emissions' each perform one MINLP solve. 'Balanced' performs N + 2 solves (2 anchors + N Pareto points)."],
        ["Number of Pareto points (N)", "2 - 20", "10",
         "Number of epsilon-constraint grid points used to trace the Pareto frontier. More points give a smoother and more informative frontier but increase total solve time approximately linearly."],
        ["Time limit per solve", "30 - 600 seconds", "300 s",
         "Maximum CPU time allocated to each individual MINLP solve. Tighter limits reduce computation time but may yield solutions with larger optimality gaps."],
        ["Relative optimality gap", "0.01 - 0.20", "5% (0.05)",
         "The solver terminates when the best feasible solution found is within this percentage of the best theoretical lower bound. A 5% gap is standard for engineering optimisation problems."],
    ],
    col_widths=[4.0, 3.5, 2.5, 5.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. TAB 5 — RESULTS
# ═══════════════════════════════════════════════════════════════
heading1("8.  Tab 5 \u2014 Results")

body(
    "The Results tab is the core output interface of ECO-FAST. The user clicks "
    "'Run Optimisation' to send all parameters from Tabs 2-4 to the "
    "FastAPI backend as a JSON payload. The backend assembles and solves the "
    "MINLP model and returns the Pareto results. Results are displayed in "
    "several sections within the same tab."
)

heading2("8.1  Pathway Configuration")
body(
    "For each Pareto point, the complete technology configuration is reported: "
    "selected conversion technology, biological pretreatment unit, mechanical "
    "pretreatment unit, HTL recovery unit (if applicable), gas upgrading unit "
    "(if applicable), and steam turbine status (INC pathway only). The "
    "associated NAC (M$ yr-1) and GHG (t CO2-eq yr-1) values are shown "
    "for each configuration."
)

heading2("8.2  Cost Breakdown")
body(
    "The total annual cost (CCTC) is decomposed into eight components for the "
    "selected pathway:"
)
add_table(
    ["Symbol", "Cost Component", "Basis of Calculation"],
    [
        ["CC_AC", "Annualised Capital Cost",
         "Equipment purchase cost scaled by the six-tenths law (exponent 0.67), multiplied by BMC = 5.4 (bare module multiplier), 1.66 (contingency/fees), and CRF."],
        ["CC_WC", "Working Capital Cost",
         "Fraction of total capital held as liquid operating reserve (typically 10-15% of CCTC)."],
        ["CC_INS", "Insurance Cost",
         "Annual insurance premium, typically 1-2% of total installed capital."],
        ["CC_UC", "Utility Cost",
         "Cost of net electricity purchased from external sources: C_elec x PW_net x T_ann."],
        ["CC_LB", "Labour Cost",
         "Direct operator labour: N_workers x C_lbr x T_ann, where N_workers is scaled by facility throughput."],
        ["CC_OC", "Overhead Cost",
         "Indirect operating costs including maintenance, administration, and supervision (typically 60-80% of labour cost)."],
        ["CC_RM", "Raw Material Cost",
         "Cost of consumables such as enzymes (ENZ pathway), chemicals, and media."],
        ["CC_DISP", "Waste Disposal Cost",
         "Cost for disposal of residue streams including non-captured char, ash, and process effluents."],
    ],
    col_widths=[2.0, 4.0, 9.0]
)

heading2("8.3  GHG Breakdown")
body(
    "Total GHG emissions (t CO2-eq yr-1) are decomposed into four terms:"
)
add_table(
    ["Symbol", "GHG Component", "Description"],
    [
        ["GHG_ind", "Indirect Emissions",
         "Scope 2 emissions from grid electricity consumed by all processing units, multiplied by the grid emission factor (phi_elec, kg CO2-eq/kWh)."],
        ["GHG_dir", "Direct Process Emissions",
         "Scope 1 emissions including process CO2 from biological degradation, fugitive CH4 (GWP = 28) and N2O (GWP = 265, IPCC AR5) from biological and thermal units, and flue gas CO2 from incineration."],
        ["GHG_aq", "Aqueous-Phase Emissions",
         "Emissions associated with treatment and disposal of aqueous by-product streams (primarily from HTL and WWT)."],
        ["GHG_disp", "Displacement Credits (subtracted)",
         "Avoided fossil emissions credited for replacing fossil products: biocrude displaces crude oil, biomethane displaces natural gas, electricity displaces grid power, compost displaces synthetic fertiliser, biosolids displace mineral amendments."],
    ],
    col_widths=[2.0, 4.0, 9.0]
)

heading2("8.4  Pareto Frontier Plot")
body(
    "For the Balanced objective mode, a cost-versus-emissions scatter plot "
    "displays all N Pareto-optimal solutions, coloured by conversion technology. "
    "The shape of the frontier reveals the magnitude of the cost penalty "
    "associated with reducing GHG emissions and identifies 'knee points' where "
    "large emission reductions can be achieved at relatively modest additional cost."
)

heading2("8.5  Results Download")
body(
    "All Pareto results - including pathway configuration, all eight cost "
    "components, all four GHG components, NAC, REV, and CCTC per Pareto point "
    "- are available for download as pareto_results.csv for further "
    "analysis in Excel, Python, or R."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. TAB 6 — ENVIRONMENTAL JUSTICE
# ═══════════════════════════════════════════════════════════════
heading1("9.  Tab 6 \u2014 Environmental Justice Assessment")

body(
    "The Environmental Justice (EJ) tab assesses the equity implications of "
    "siting a food waste facility at the specified location by combining "
    "optimisation results with community demographic vulnerability data. "
    "The tab is active only when a valid zip code has been entered in Tab 2 "
    "and an optimisation has been completed in Tab 5. The methodology follows "
    "the EPA EJScreen framework (US EPA, 2023)."
)

heading2("9.1  Facility Location and Buffer Zones")
body(
    "An interactive map rendered with Folium displays the facility location "
    "with three concentric buffer zones:"
)
add_table(
    ["Buffer Zone", "Radius", "Colour", "Interpretation"],
    [
        ["Immediate impact zone", "1 km", "Red",
         "Residents within this radius experience the most direct exposure to facility emissions, traffic, and noise."],
        ["Surrounding neighbourhood", "3 km", "Orange",
         "The area most commonly affected by atmospheric dispersion of air pollutants under typical meteorological conditions."],
        ["Broader community", "5 km", "Green",
         "The wider area potentially affected by cumulative environmental burdens from facility operations."],
    ],
    col_widths=[4.0, 2.0, 2.0, 7.0]
)
body(
    "Note: the buffer rings are illustrative. Demographic statistics are reported "
    "at the zip code tabulation area (ZCTA) level from the US Census ACS 5-year "
    "estimates - not computed separately for each buffer radius."
)

heading2("9.2  Community Demographics")
body(
    "Demographic data for the facility zip code are retrieved from the US Census "
    "Bureau ACS 5-year estimates (2023 vintage, accessed via the Census Data API). "
    "Four indicators are displayed:"
)
add_table(
    ["Indicator", "Formula", "National Average", "Interpretation"],
    [
        ["% People of Colour (%PoC)",
         "(Total pop. - Non-Hispanic white) / Total pop. x 100",
         "40.0% (ACS 2022)",
         "Communities with %PoC above the national average are considered more demographically vulnerable under EJScreen."],
        ["% Low-Income (%LI)",
         "Population below federal poverty line / Poverty universe x 100",
         "29.0% (ACS 2022)",
         "Reflects economic vulnerability. Low-income communities often bear disproportionate environmental burdens."],
        ["Demographic Index (DI)",
         "(%PoC + %LI) / 2",
         "34.5%",
         "Composite vulnerability indicator used as the EJScreen Demographic Index."],
        ["Total Population",
         "ACS B01003_001E",
         "-",
         "Total resident population of the zip code tabulation area; used as a multiplier in the EJ Index formula."],
    ],
    col_widths=[3.5, 4.5, 3.0, 4.0]
)

heading2("9.3  EJ Index Calculation")
body(
    "An EJ Index is computed for each conversion technology using the minimum "
    "GHG emissions achieved by that technology across all Pareto solutions. "
    "The index formula is:"
)
body(
    "EJ Index = GHG (t CO2-eq yr-1)  x  (Local DI - National Average DI)  x  Total Population",
    indent=1
)
body(
    "A higher EJ Index indicates greater cumulative environmental burden placed "
    "on a more vulnerable community. Pathways are ranked from lowest to highest "
    "EJ Index and classified into three concern levels (Low, Moderate, High) "
    "based on their relative position within the range of scores across all "
    "evaluated pathways. A negative EJ Index arises when the local DI is below "
    "the national average, indicating the facility would be sited in a "
    "relatively less vulnerable community."
)

heading2("9.4  Recommendation")
body(
    "A plain-language recommendation section identifies the technology pathway "
    "with the lowest EJ Index and quantifies the reduction in community burden "
    "relative to the worst-performing pathway, expressed in absolute EJ Index "
    "units. This output is intended to support environmental justice reporting "
    "in permit applications, environmental impact assessments, and regulatory "
    "submissions."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. MATHEMATICAL MODEL SUMMARY
# ═══════════════════════════════════════════════════════════════
heading1("10.  Mathematical Model Summary")

body(
    "ECO-FAST solves a Mixed-Integer Nonlinear Program (MINLP) comprising binary "
    "technology selection variables, continuous mass flow and economic variables, "
    "and a set of linear and nonlinear constraints. The complete mathematical "
    "formulation with all numbered equations is presented in the Methods section "
    "of the associated publication. Key model statistics are summarised below."
)

add_table(
    ["Attribute", "Value / Description"],
    [
        ["Model type", "MINLP (Mixed-Integer Nonlinear Program)"],
        ["Number of material streams", "61"],
        ["Number of stream components", "8 (WATER, CBH, PRT, FAT, OTH, ASH, FC, ENZYME)"],
        ["Number of processing units", "15"],
        ["Binary decision variables", "~20 (technology selection and routing switches)"],
        ["Continuous variables", "~500+ (stream flows, capacities, costs, emissions, sizing variables)"],
        ["Equality constraints", "~400 (mass balances, sizing, cost, GHG component equations)"],
        ["Nonlinear terms", "Capital cost scaling law (Qc/Q0)^0.67 per unit; steam turbine enthalpy calculation"],
        ["Objective functions", "2: NAC (M$ yr-1) and GHG (t CO2-eq yr-1)"],
        ["Multi-objective method", "Epsilon-constraint Pareto frontier (N user-defined points)"],
        ["Solver", "SBB (Simple Branch and Bound) via GAMSPy v1.23.1"],
        ["NLP subproblem solver", "CONOPT (local NLP; called at each B&B node)"],
        ["Optimality", "Local optimum (SBB is a local MINLP solver; global optimality not guaranteed for nonconvex problems)"],
        ["Relative optimality gap", "5% (default; user-adjustable)"],
        ["Time limit per solve", "300 seconds (default; user-adjustable)"],
    ],
    col_widths=[5.5, 9.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. DATA FLOW BETWEEN TABS
# ═══════════════════════════════════════════════════════════════
heading1("11.  Data Flow Between Tabs")

body(
    "All user inputs are held in Streamlit session state throughout a session. "
    "When the user clicks 'Run Optimisation' in Tab 5, a unified input dictionary "
    "is assembled from all session state values and sent to the FastAPI backend "
    "as a single JSON request. The data flow is strictly sequential across tabs."
)

add_table(
    ["Source Tab", "Data Produced", "Used In"],
    [
        ["Tab 2 - Feed Inputs",
         "Q_f (kg/hr), T_ann (hr/yr), wet-basis composition (7 fractions), zip code",
         "Backend model - all process units use feed flow and composition as boundary conditions. Zip code used in Tab 6 for EJ assessment."],
        ["Tab 3 - Technology Specifications",
         "Unit-specific engineering parameters for all 15 technologies (yields, efficiencies, HRTs, emission factors)",
         "Backend model - overrides default parameter values in parameters.py for the selected technology parameters."],
        ["Tab 4 - Cost Specifications",
         "C_elec, C_tip, C_lbr, discount rate r, plant lifetime n, product prices, objective mode, N Pareto points, time limit, gap",
         "Backend model - costing.py (cost equations), solver.py (objective and epsilon grid configuration)."],
        ["Tab 5 - Results",
         "Pareto dataframe: NAC, GHG, all cost and GHG components, pathway configuration per Pareto point",
         "Displayed in Tab 5. GHG values passed to Tab 6 for EJ Index calculation. Full dataframe available as CSV download."],
        ["Tab 6 - EJ Assessment",
         "EJ Index per pathway, community demographics, facility map",
         "Final output only. Not fed back to any upstream tab."],
    ],
    col_widths=[3.5, 6.0, 5.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. SOFTWARE STACK AND DEPLOYMENT
# ═══════════════════════════════════════════════════════════════
heading1("12.  Software Stack and Deployment")

heading2("12.1  Software Dependencies")
add_table(
    ["Package", "Version", "Role"],
    [
        ["Python", "3.12", "Runtime environment for all model and application code"],
        ["GAMSPy", "1.23.1", "Python-native algebraic modelling interface to the GAMS solver system"],
        ["GAMS / SBB", "System license", "Branch-and-bound MINLP solver; SBB is the local MINLP solver used for all optimisation runs"],
        ["Streamlit", "Latest stable", "Frontend web application framework; renders all six tabs and interactive widgets"],
        ["FastAPI", "Latest stable", "REST API backend framework; receives optimisation requests and returns results as JSON"],
        ["Uvicorn", "Latest stable", "ASGI web server that hosts the FastAPI application"],
        ["Pandas", "Latest stable", "Tabular data handling, Pareto result formatting, and CSV export"],
        ["Matplotlib", "Latest stable", "Composition donut chart and Pareto frontier scatter plot in the Results tab"],
        ["Folium", "Latest stable", "Interactive facility location map with buffer zones in the EJ tab"],
        ["streamlit-folium", "Latest stable", "Streamlit component for rendering Folium maps within the Streamlit interface"],
        ["Requests", "Latest stable", "HTTP requests to the Zippopotam.us geocoding API and US Census Bureau ACS data API"],
        ["python-docx", "Latest stable", "Used to generate this report document programmatically"],
    ],
    col_widths=[3.5, 2.5, 9.0]
)

heading2("12.2  Starting the Application")
body(
    "Both processes must be started in separate terminal windows. The backend "
    "API must be started first, followed by the Streamlit frontend."
)
bullet("Terminal 1 - Backend API:  ",
    "py -3.12 -m uvicorn api:app --host 127.0.0.1 --port 8000")
bullet("Terminal 2 - Frontend:  ",
    "py -3.12 -m streamlit run app.py")
body(
    "Once both processes are running, the application is accessible in a web "
    "browser at http://localhost:8501. The API health endpoint can be verified "
    "at http://127.0.0.1:8000/health, which returns {\"status\": \"ok\"} when "
    "the backend is operating correctly."
)

heading2("12.3  Sharing with External Stakeholders (Temporary Access)")
body(
    "For sharing the tool with external collaborators or regulatory agencies "
    "(e.g. EPA) during a meeting or review period, the application can be "
    "temporarily exposed to the internet using the ngrok tunnelling service:"
)
bullet("Step 1:  ", "Start both backend and frontend processes as described above.")
bullet("Step 2:  ", "Run: ngrok http 8501")
bullet("Step 3:  ",
    "Share the public URL provided by ngrok (e.g. https://abc123.ngrok-free.app) "
    "with stakeholders. All computation remains on the host machine; the "
    "GAMS license remains valid on the local installation.")
body(
    "This approach requires no cloud infrastructure and no license transfer. "
    "The ngrok free tier displays a one-time warning page before the application "
    "loads; a paid ngrok account removes this warning."
)

heading2("12.4  Permanent Cloud Deployment")
body(
    "For permanent hosting accessible to all stakeholders at any time, the "
    "application can be deployed on a cloud virtual machine (e.g. Google Cloud "
    "Compute Engine, AWS EC2, or Azure VM). Both backend and frontend processes "
    "run on the same VM, and a public IP address or custom domain provides "
    "permanent access."
)
body(
    "The key requirement for cloud deployment is a valid GAMS license installed "
    "on the cloud VM. Standard GAMS node-locked licenses may be transferred to "
    "a cloud VM (up to three transfers per license term under current GAMS EULA "
    "terms). Cloud deployment costs are approximately $25-50 per month for a "
    "general-purpose VM instance of appropriate size (e.g. 2 vCPUs, 8 GB RAM)."
)

# ── Save ──────────────────────────────────────────────────────
out = r"C:\Users\harri\OneDrive\Desktop\foodwaste_optimization\ECO-FAST_Tool_Description_Report.docx"
doc.save(out)
print(f"\nDone. Report saved to:\n{out}")
